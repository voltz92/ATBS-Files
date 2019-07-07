# word docs

# each word document has the following objects ::
# Document = represents the whole document itself
# Paragraph = represents the paragraph objects
# Run = represents a string of characters with similar styles. everytime the style changes a new run is created

import docx
import os
os.chdir('D:\\Files\\Git Repos\\ATBS-Files\\Office Files\\Word-Files')

doc = docx.Document('demo.docx')


for i in doc.paragraphs:
    print(i.text)

p = doc.paragraphs[1]
# we can change the style of the paragraphs by using the style attribute
p.style = 'Title'
doc.save('demo2.docx')

# each paragraphs has runs.
print(p.runs)

# the following code turns all runs in the 2nd paragraph of the document to bold.
for i in p.runs:
    i.bold = True
    i.underline = True

doc.save('demo3.docx')


