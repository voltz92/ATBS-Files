# creating a new doc 

import docx
import os

os.chdir('D:\\Files\\Git Repos\\ATBS-Files\\Office Files\\Word-Files')

# this creates a new doc object
d = docx.Document()

# each document has an .add_paragraph() method. the arguments of this method are added as text
d.add_paragraph('This is my title')
# sets the paragraph style to Title
d.paragraphs[0].style = 'Title'
# adding a sencond para
d.add_paragraph()
# assigning it to a variable
p = d.paragraphs[1]

# adding a new run, which has the bold style.
p.add_run('This is a sentence which is bold. ').bold = True
p.add_run('This sentence has an underline.').underline = True
# saving the file
d.save('python_made.docx')

